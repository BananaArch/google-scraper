from bs4 import BeautifulSoup as bs
import requests
from docx import Document as doc
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches

# final vars
url_prefix = "https://www.google.com/search?q="
url_suffix = " definition"
file_name = 'terms'
title = "TERMS TEST"
name = "Dwayne \"THE ROCK\" JOHNSON!!!"
font_name = 'Arial'
color1 = RGBColor(255, 101, 66)
color2 = RGBColor(255, 144, 37)
color3 = RGBColor(255, 186, 8)
color4 = RGBColor(62, 102, 128)

headers = {
    'Access-Control-Allow-Origin': '*',
    'Access-Control-Allow-Methods': 'GET',
    'Access-Control-Allow-Headers': 'Content-Type',
    'Access-Control-Max-Age': '3600',
    'User-Agent': 'Mozilla/5.0 (X11; Ubuntu; Linux x86_64; rv:52.0) Gecko/20100101 Firefox/52.0'
    }


## Creates Document

document = doc()

## Margins
# sections = document.sections
# for section in sections:
#     section.top_margin = Inches(1)
#     section.bottom_margin = Inches(1)
#     section.left_margin = Inches(1)
#     section.right_margin = Inches(1)

## Style 

h1 = document.styles['Heading 1']
h1.font.name = font_name
h1.font.size = Pt(16)
h1.font.bold = True
h1.font.italic = False
h1.font.color.rgb = color1

h2 = document.styles['Heading 2']
h2.font.name = font_name
h2.font.size = Pt(12)
h2.font.bold = True
h2.font.italic = False
h2.font.color.rgb = color2

h3 = document.styles['Heading 3']
h3.font.name = font_name
h3.font.size = Pt(12)
h3.font.bold = True
h3.font.italic = True
h3.font.underline = True
h3.font.color.rgb = color3

h4 = document.styles['Normal']
h4.font.name = font_name
h4.font.size = Pt(12)
h4.font.italic = True
h4.font.color.rgb = RGBColor(0, 0, 0)

bt = document.styles['Body Text']
bt.font.name = font_name
bt.font.size = Pt(12)
bt.font.italic = False
bt.font.color.rgb = color4

## Write Name

pname = document.add_paragraph(name)
pname.alignment = WD_ALIGN_PARAGRAPH.RIGHT
pname.style = h4

## Write title

ptitle = document.add_paragraph(title)
ptitle.style = h1


## Load terms.txt

file = open('{}.txt'.format(file_name), 'r')
f = file.readlines()

## Saves terms in array

terms = []
for line in f:
    terms.append(line.strip().upper())


## Searches all terms

for term in terms:

    ## Check if Section is title
    if (term[0:2].lower() == '\\t'):

        psectitle = document.add_paragraph(term[2:])
        psectitle.style = h2
        print(term)
        
        continue

    ## Load webpage
    
    req = requests.get(url_prefix + term + " " + url_suffix, headers)
    soup = bs(req.text, 'html.parser')
        
    ## Get Definition
    
    definition = soup.find(class_ = 'BNeawe s3v9rd AP7Wnd').text.strip().capitalize()
    
    # Add Term to Doc

    pterm = document.add_paragraph(term)
    pterm.style = h3

    pdef = document.add_paragraph(definition)
    pdef.style = bt

    print(term)
    print(definition + "\n\n\n")

## Save Document

document.save('{}.docx'.format(file_name))